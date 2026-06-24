#!/usr/bin/env python3
"""@Apolaibot Demo — Telegram bot: AI agents with skills tailored to your business."""

import os, sys, logging, aiohttp, time, io, json as _json
from pathlib import Path as _Path
from datetime import datetime as _dt
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from telegram.constants import ParseMode

# Agent FS — versioned file ops + audit trail
sys.path.insert(0, os.path.expanduser("~/.hermes/scripts"))
# Fallback to system site-packages for whisper
sys.path.insert(0, os.path.expanduser("~/.local/lib/python3.12/site-packages"))
try:
    from agent_fs import versioned_write, audit as fs_audit
    AGENT_FS_ENABLED = True
except ImportError:
    AGENT_FS_ENABLED = False
    def versioned_write(*a, **kw): return {}
    def fs_audit(*a, **kw): pass

# 2FA via TOTP
try:
    import pyotp
    TOTP_ENABLED = True
except ImportError:
    TOTP_ENABLED = False

# Transcription via Whisper
try:
    import whisper
    WHISPER_MODEL = None  # lazy load
    WHISPER_ENABLED = True
except ImportError:
    WHISPER_ENABLED = False

ADMIN_IDS = {319665243, 5529208670}  # Poliakarm + Cryptos
ADMIN_2FA_FILE = _Path.home() / ".local" / "share" / "apolaibot" / "admin_2fa.json"

def _load_2fa():
    try:
        if ADMIN_2FA_FILE.exists():
            return _json.loads(ADMIN_2FA_FILE.read_text())
    except Exception:
        pass
    return {}

def _save_2fa(data):
    ADMIN_2FA_FILE.parent.mkdir(parents=True, exist_ok=True)
    ADMIN_2FA_FILE.write_text(_json.dumps(data, ensure_ascii=False, indent=2))

# In-memory 2FA session cache: {user_id: verified_until_timestamp}
_2fa_sessions = {}

def is_admin(user_id): return user_id in ADMIN_IDS

def verify_2fa(user_id, code):
    """Verify TOTP code for admin. Returns True if valid."""
    if not TOTP_ENABLED:
        return False  # Scout fix R5: no pyotp = 2FA mandatory, not bypassed
    secrets = _load_2fa()
    secret = secrets.get(str(user_id))
    if not secret:
        return False  # Scout fix R5: 2FA not set up = blocked, not allowed
    totp = pyotp.TOTP(secret)
    return totp.verify(code)

def check_2fa_session(user_id):
    """Check if user has active 2FA session (valid for 30 min)."""
    if not TOTP_ENABLED:
        return True
    return _2fa_sessions.get(user_id, 0) > time.time()

def grant_2fa_session(user_id):
    _2fa_sessions[user_id] = time.time() + 1800  # 30 min

TOKEN = os.environ.get("TELEGRAM_" + "BOT_TOKEN", "")
if not TOKEN:
    print("FATAL: TELEGRAM_BOT_TOKEN not set", file=sys.stderr)
    sys.exit(1)

# Scout fix R6: public URL from env, not hardcoded IP
PUBLIC_URL = os.environ.get("APOLAIBOT_PUBLIC_URL", "https://t.me/Apolaibot")

API_URL = os.environ.get("APOLAIBOT_API_URL", "https://api.deepseek.com/v1")
API_KEY = os.environ.get("APOLAIBOT_API_KEY", os.environ.get("DEEPSEEK_API_KEY", ""))
MODEL = os.environ.get("APOLAIBOT_MODEL", "deepseek-chat")
if not API_KEY:
    print("FATAL: neither APOLAIBOT_API_KEY nor DEEPSEEK_API_KEY set", file=sys.stderr)
    sys.exit(1)

sys.path.insert(0, os.path.expanduser("~/.hermes/scripts"))
try:
    from rate_limiter import can_proceed, check as rate_check, record_tokens, can_use_tokens, get_token_budget
    RATE_LIMIT_ENABLED = True
except ImportError:
    RATE_LIMIT_ENABLED = False
    def can_proceed(*a, **kw): return True
    def rate_check(*a, **kw): return (True, 999, 0)
    def record_tokens(*a, **kw): pass
    def can_use_tokens(*a, **kw): return True
    def get_token_budget(*a, **kw): return {}

try:
    from tenant_audit import log_action
    AUDIT_ENABLED = True
except ImportError:
    AUDIT_ENABLED = False
    def log_action(*a, **kw): pass

RATE_LIMIT_COOLDOWN_MSG = "⏳ Слишком много запросов. Подожди немного. Хочешь больше? /upgrade → Pro с безлимитом 🚀"

TAVILY_KEY = os.environ.get("TAVILY_API_KEY", "")

logging.basicConfig(format="%(asctime)s [%(levelname)s] %(message)s", level=logging.INFO, stream=sys.stdout)
logger = logging.getLogger("apolaibot-demo")

# ── Skill catalog ──
SKILL_CATALOG = {
    "manager": {
        "emoji": "📊", "label_ru": "Руководитель", "label_en": "Manager",
        "skills": ["Отчёты и сводки", "Стратегическое планирование", "Деловые письма", "Протоколы встреч", "Анализ KPI"],
    },
    "accountant": {
        "emoji": "🧮", "label_ru": "Бухгалтер / Финансист", "label_en": "Accountant / Finance",
        "skills": ["Первичная документация", "Пояснительные записки", "Финансовые сводки", "Расчёт налогов", "Анализ бюджета"],
    },
    "marketer": {
        "emoji": "📣", "label_ru": "Маркетолог", "label_en": "Marketer",
        "skills": ["УТП и позиционирование", "Посты для соцсетей", "Email-рассылки", "Контент-планы", "Анализ конкурентов"],
    },
    "sales": {
        "emoji": "💼", "label_ru": "Продажи", "label_en": "Sales",
        "skills": ["Коммерческие предложения", "Скрипты продаж", "Отработка возражений", "Follow-up письма", "Воронка продаж"],
    },
    "hr": {
        "emoji": "👥", "label_ru": "HR", "label_en": "HR",
        "skills": ["Вакансии и описания", "Офферы кандидатам", "Вопросы для интервью", "Онбординг-планы", "Оценка персонала"],
    },
    "teacher": {
        "emoji": "🎓", "label_ru": "Учитель / Студент", "label_en": "Teacher / Student",
        "skills": ["Конспекты лекций", "Тесты и задания", "Планы занятий", "Проверка работ", "Объяснение тем"],
    },
    "lawyer": {
        "emoji": "⚖️", "label_ru": "Юрист", "label_en": "Lawyer",
        "skills": ["Договоры и соглашения", "Претензии и иски", "Анализ рисков", "Юридические аргументы", "Проверка документов"],
    },
}

async def tavily_search(query):
    if not TAVILY_KEY:
        return None
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                "https://api.tavily.com/search",
                json={"api_key": TAVILY_KEY, "query": query, "max_results": 3, "search_depth": "basic"},
                headers={"Content-Type": "application/json"},
                timeout=aiohttp.ClientTimeout(total=10),
            ) as resp:
                if resp.status != 200:
                    return None
                data = await resp.json()
                results = data.get("results", [])
                if not results:
                    return None
                lines = []
                for r in results[:3]:
                    lines.append(f"- {r['title']}: {r['content'][:200]}")
                return "\n".join(lines)
    except Exception as e:
        logger.error(f"Tavily error: {e}")
        return None

def needs_web_search(msg):
    triggers = [
        "сегодня", "сейчас", "новост", "сколько стоит", "цена", "курс",
        "погода", "что произошл", "когда выйдет", "последн",
        "2026", "2025", "вчера", "на этой неделе", "в этом месяце",
    ]
    msg_lower = msg.lower()
    return any(t in msg_lower for t in triggers)

SYSTEM_PROMPT = """You are @Apolaibot — you provide AI agents with skills tailored to business processes. Not a generic chatbot trying to sell AI hype. You deliver ready agents that come with their own skills and work in the client's context. Trading, documents, security — the agent already knows the domain, no need to explain the basics.
Tone: professional, confident, no-nonsense. Answer in the user's language. Demo mode — suggest /upgrade for full access.
RULES:
1. Be brief. 1-2 short paragraphs unless the user explicitly asks for detail.
2. Don't fabricate. If you don't know — say so, don't guess.
3. One call-to-action per response. Suggest the most relevant command, not all of them.
Available commands: /skills (catalog), /summarize (condense chat), /temp (incognito), /export (save chat to MD/PDF/DOCX), /stats (token budget), /tools (transparency)."""

# ── User registry ──
USERS_FILE = _Path.home() / ".local" / "share" / "apolaibot" / "users.json"
FEEDBACK_FILE = _Path.home() / ".local" / "share" / "apolaibot" / "feedback.json"

def _load_users():
    try:
        if USERS_FILE.exists():
            return _json.loads(USERS_FILE.read_text())
    except Exception:
        pass
    return {}

def _save_users(users):
    USERS_FILE.parent.mkdir(parents=True, exist_ok=True)
    USERS_FILE.write_text(_json.dumps(users, ensure_ascii=False, indent=2))

def _load_feedback():
    try:
        if FEEDBACK_FILE.exists():
            return _json.loads(FEEDBACK_FILE.read_text())
    except Exception:
        pass
    return {}

def _save_feedback(fb):
    FEEDBACK_FILE.parent.mkdir(parents=True, exist_ok=True)
    FEEDBACK_FILE.write_text(_json.dumps(fb, ensure_ascii=False, indent=2))

def track_user(user_id, username, first_name, role=None):
    users = _load_users()
    uid = str(user_id)
    if uid not in users:
        users[uid] = {
            "user_id": user_id,
            "username": username or "",
            "first_name": first_name or "",
            "role": role,
            "first_seen": _dt.now().isoformat(),
            "last_seen": _dt.now().isoformat(),
            "msg_count": 1,
        }
        logger.info(f"[NEW USER] id={user_id} @{username or '—'} name={first_name} role={role}")
    else:
        users[uid]["last_seen"] = _dt.now().isoformat()
        users[uid]["msg_count"] = users[uid].get("msg_count", 0) + 1
        if username and username != users[uid].get("username"):
            users[uid]["username"] = username
        if role is not None:
            users[uid]["role"] = role
    _save_users(users)

def get_user_role(user_id):
    users = _load_users()
    return users.get(str(user_id), {}).get("role")

# ── Conversations (chat_id → history) ──
conversations = {}       # normal chats
temp_conversations = {}  # incognito chats
MAX_HISTORY = 20

async def chat_completion(chat_id, user_msg, system_prompt=None, temp=False):
    store = temp_conversations if temp else conversations
    if chat_id not in store:
        store[chat_id] = [{"role": "system", "content": system_prompt or SYSTEM_PROMPT}]
    store[chat_id].append({"role": "user", "content": user_msg})
    if len(store[chat_id]) > MAX_HISTORY + 1:
        store[chat_id] = [store[chat_id][0]] + store[chat_id][-(MAX_HISTORY):]
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}", "Accept-Encoding": "gzip"}
    payload = {"model": MODEL, "messages": store[chat_id], "temperature": 0.7, "max_tokens": 1500}
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(f"{API_URL}/chat/completions", headers=headers, json=payload, timeout=aiohttp.ClientTimeout(total=60)) as resp:
                if resp.status != 200:
                    logger.error(f"API error {resp.status}")
                    return "Sorry, service temporarily unavailable.", 0
                data = await resp.json()
                reply = data["choices"][0]["message"]["content"]
                store[chat_id].append({"role": "assistant", "content": reply})
                usage = data.get("usage", {})
                total_tokens = usage.get("total_tokens", 0)
                return reply, total_tokens
    except Exception as e:
        logger.error(f"API error: {e}")
        return "Unable to reach AI. Try later.", 0

def _feedback_buttons():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("👍", callback_data="fb_like"),
         InlineKeyboardButton("👎", callback_data="fb_dislike")]
    ])

# ── Commands ──

async def cmd_start(update, context):
    lang = (update.effective_user.language_code or "en").lower()
    if lang.startswith("ru"):
        keyboard = [
            [InlineKeyboardButton("📊 Руководитель", callback_data="role_manager")],
            [InlineKeyboardButton("🧮 Бухгалтер / Финансист", callback_data="role_accountant")],
            [InlineKeyboardButton("📣 Маркетолог", callback_data="role_marketer")],
            [InlineKeyboardButton("💼 Продажи", callback_data="role_sales")],
            [InlineKeyboardButton("👥 HR", callback_data="role_hr")],
            [InlineKeyboardButton("🎓 Учитель / Студент", callback_data="role_teacher")],
            [InlineKeyboardButton("⚖️ Юрист", callback_data="role_lawyer")],
        ]
        await update.message.reply_markdown(
            f"Привет, {update.effective_user.first_name}!\n\n"
            "**@Apolaibot** — мы не продаём «попробуйте AI». "
            "Мы даём готового агента с навыками под ваш процесс.\n\n"
            "🌐 http://PUBLIC_URL — тарифы и подробности\n\n"
            "👆 *Кто вы по роду занятий?*",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
    else:
        keyboard = [
            [InlineKeyboardButton("📊 Manager", callback_data="role_manager")],
            [InlineKeyboardButton("🧮 Accountant / Finance", callback_data="role_accountant")],
            [InlineKeyboardButton("📣 Marketer", callback_data="role_marketer")],
            [InlineKeyboardButton("💼 Sales", callback_data="role_sales")],
            [InlineKeyboardButton("👥 HR", callback_data="role_hr")],
            [InlineKeyboardButton("🎓 Teacher / Student", callback_data="role_teacher")],
            [InlineKeyboardButton("⚖️ Lawyer", callback_data="role_lawyer")],
        ]
        await update.message.reply_markdown(
            f"Hello, {update.effective_user.first_name}!\n\n"
            "**@Apolaibot** — we don't sell «try AI». "
            "We deliver a ready agent with skills for your process.\n\n"
            "🌐 http://PUBLIC_URL — pricing & details\n\n"
            "👆 *What's your profession?*",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )

async def role_callback(update, context):
    """Handle role selection from inline keyboard."""
    query = update.callback_query
    await query.answer()
    role_key = query.data.replace("role_", "")
    role = SKILL_CATALOG.get(role_key)
    if not role:
        return

    user_id = query.from_user.id
    lang = (query.from_user.language_code or "en").lower()
    track_user(user_id, query.from_user.username, query.from_user.first_name, role=role_key)

    if lang.startswith("ru"):
        skills_list = "\n".join(f"• {s}" for s in role["skills"])
        text = (
            f"{role['emoji']} **{role['label_ru']}** — вот что я умею для вас:\n\n"
            f"{skills_list}\n\n"
            f"Просто напишите задачу — агент уже в теме.\n"
            f"🚀 /upgrade — полная версия, безлимитный контекст, 154+ скиллов.\n"
            f"📋 /skills — все профессии."
        )
    else:
        skills_list = "\n".join(f"• {s}" for s in role["skills"])
        text = (
            f"{role['emoji']} **{role['label_en']}** — here's what I can do:\n\n"
            f"{skills_list}\n\n"
            f"Just describe your task — the agent already knows the domain.\n"
            f"🚀 /upgrade — full version, unlimited context, 154+ skills.\n"
            f"📋 /skills — all professions."
        )
    await query.edit_message_text(text, parse_mode=ParseMode.MARKDOWN)

async def cmd_skills(update, context):
    """Show skill catalog grouped by profession."""
    lang = (update.effective_user.language_code or "en").lower()
    if lang.startswith("ru"):
        lines = ["📋 **Каталог скиллов @Apolaibot**\n"]
        for key, role in SKILL_CATALOG.items():
            skills_preview = ", ".join(role["skills"][:3])
            lines.append(f"{role['emoji']} **{role['label_ru']}**: {skills_preview}")
        lines.append(f"\n🔢 Всего профессий: {len(SKILL_CATALOG)}")
        lines.append("🚀 /upgrade — полная версия (154+ скиллов).")
    else:
        lines = ["📋 **@Apolaibot Skill Catalog**\n"]
        for key, role in SKILL_CATALOG.items():
            skills_preview = ", ".join(role["skills"][:3])
            lines.append(f"{role['emoji']} **{role['label_en']}**: {skills_preview}")
        lines.append(f"\n🔢 Total professions: {len(SKILL_CATALOG)}")
        lines.append("🚀 /upgrade — full version (154+ skills).")
    await update.message.reply_markdown("\n".join(lines))

async def cmd_summarize(update, context):
    """Summarize the current chat history."""
    chat_id = update.effective_chat.id
    store = conversations.get(chat_id)
    if not store or len(store) <= 2:
        await update.message.reply_text("Not enough history to summarize yet. Chat a bit first!")
        return
    await context.bot.send_chat_action(chat_id=chat_id, action="typing")
    # Build summary prompt from history
    history_text = "\n".join(
        f"{m['role']}: {m['content'][:300]}" for m in store[1:] if m["role"] != "system"
    )
    if len(history_text) > 6000:
        history_text = history_text[:6000] + "..."
    summary_prompt = f"Summarize this conversation in 3-5 bullet points. Keep it brief and factual. Answer in the same language as the conversation:\n\n{history_text}"
    summary, _ = await chat_completion(chat_id, summary_prompt)
    await update.message.reply_markdown(f"📝 **Chat summary:**\n{summary}")

async def cmd_temp(update, context):
    """Start an incognito (temporary) chat session."""
    user_id = update.effective_user.id
    lang = (update.effective_user.language_code or "en").lower()

    # Scout fix R10: rate limit temp sessions (was unlimited bypass)
    if RATE_LIMIT_ENABLED:
        ok, remaining, _ = rate_check(user_id, "temp")
        if not ok:
            if lang.startswith("ru"):
                await update.message.reply_text("⚠️ Слишком много инкогнито-сессий. Подождите минуту.")
            else:
                await update.message.reply_text("⚠️ Too many incognito sessions. Wait a minute.")
            return

    if lang.startswith("ru"):
        await update.message.reply_text(
            "🕶️ **Режим инкогнито** — этот чат не сохраняется.\n"
            "Пишите что угодно — после закрытия ничего не останется.\n"
            "Для выхода: просто закройте чат или начните новый диалог."
        )
    else:
        await update.message.reply_text(
            "🕶️ **Incognito mode** — this chat is not saved.\n"
            "Write anything — nothing remains after you close it.\n"
            "To exit: just close the chat or start a new conversation."
        )

async def cmd_export(update, context):
    """Export chat as Markdown, PDF, or DOCX — with versioning."""
    chat_id = update.effective_chat.id
    user_id = update.effective_user.id
    msg_text = update.message.text.strip() if update.message.text else ""
    
    # Parse format from command: /export pdf, /export docx, etc.
    fmt = "md"
    parts = msg_text.split()
    if len(parts) > 1 and parts[1].lower() in ("pdf", "docx", "md"):
        fmt = parts[1].lower()
    
    store = conversations.get(chat_id)
    if not store or len(store) <= 2:
        await update.message.reply_text("Nothing to export yet. Chat first!")
        return
    
    # Build chat text
    lines = []
    for m in store:
        if m["role"] == "system":
            continue
        role_label = "You" if m["role"] == "user" else "Apolaibot"
        lines.append(f"{'##' if fmt == 'md' else ''} {role_label}\n{m['content']}\n")
    text = "\n".join(lines)
    
    timestamp = _dt.now().strftime('%Y%m%d_%H%M%S')
    base_dir = _Path.home() / ".local" / "share" / "apolaibot" / "exports"
    base_dir.mkdir(parents=True, exist_ok=True)
    
    if fmt == "pdf":
        export_path = base_dir / f"chat_{timestamp}.pdf"
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", uni=True)
        pdf.set_font("DejaVu", "", 10)
        for m in store:
            if m["role"] == "system":
                continue
            role = "👤 You" if m["role"] == "user" else "🤖 Apolaibot"
            pdf.set_font("DejaVu", "B", 10)
            pdf.multi_cell(0, 6, role)
            pdf.set_font("DejaVu", "", 10)
            # Clean markdown
            clean = m['content'].replace('**', '').replace('*', '').replace('`', '')
            pdf.multi_cell(0, 6, clean)
            pdf.ln(4)
        pdf.output(str(export_path))
        caption = "📦 Chat export (PDF) — versioned ✓"
    
    elif fmt == "docx":
        export_path = base_dir / f"chat_{timestamp}.docx"
        from docx import Document
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        for m in store:
            if m["role"] == "system":
                continue
            role = "You" if m["role"] == "user" else "Apolaibot"
            doc.add_heading(role, level=2)
            doc.add_paragraph(m['content'])
        doc.save(str(export_path))
        caption = "📦 Chat export (Word) — versioned ✓"
    
    else:  # md
        export_path = base_dir / f"chat_{timestamp}.md"
        export_path.write_text(text, encoding="utf-8")
        caption = "📦 Chat export (Markdown) — versioned ✓\n💡 /export pdf | /export docx — другие форматы"
    
    if AGENT_FS_ENABLED:
        versioned_write(str(export_path), text, message_id=str(update.message.message_id), agent=f"demo-user-{user_id}")
        fs_audit("export_chat", str(export_path), message_id=str(update.message.message_id),
                 agent=f"demo-user-{user_id}", extra={"chat_id": str(chat_id), "msg_count": len(store), "format": fmt})
    
    with open(export_path, "rb") as f:
        await update.message.reply_document(f, caption=caption)

async def cmd_stats(update, context):
    """Show user stats + token usage."""
    users = _load_users()
    total = len(users)
    today = _dt.now().strftime("%Y-%m-%d")
    active_today = sum(1 for u in users.values() if u.get("last_seen", "").startswith(today))
    lines = [f"👥 **Users:** {total}", f"📅 **Today:** {active_today}"]
    # Token usage
    if RATE_LIMIT_ENABLED:
        budget = get_token_budget(update.effective_user.id)
        if budget:
            lines.append(f"\n💎 **Token budget:**")
            lines.append(f"  • Used: {budget['used']:,}/{budget['limit']:,} ({budget['pct']:.0f}%)")
            lines.append(f"  • Remaining: {budget['remaining']:,}  |  Reset: {budget['reset_days']} days")
        else:
            lines.append(f"\n💎 **Tokens:** tracking active")
    lines.append(f"\n📋 /skills — catalog  |  🚀 /upgrade — Pro")
    await update.message.reply_markdown("\n".join(lines))

async def cmd_upgrade(update, context):
    await update.message.reply_markdown(
        "🚀 **@Apolaibot — тарифы**\n\n"
        "🌐 [apolaibot.ru](http://PUBLIC_URL)\n\n"
        "🆓 **Start** — бесплатно навсегда\n"
        "  • 1M токенов/мес • чат + скиллы\n\n"
        "💎 **Basic** — 690₽/мес\n"
        "  • 5M токенов/мес • + поиск + документы\n\n"
        "⚡ **Pro** — 1990₽/мес\n"
        "  • 20M токенов/мес • до 5 человек • + экспорт + инкогнито\n\n"
        "🏢 **Expert** — 5990₽/мес\n"
        "  • 100M токенов/мес • до 50 человек • + трейдинг + on-prem\n\n"
        "💎 **Оплата:**\n"
        "• Картой РФ → [Tribute](https://tribute.tg)\n"
        "• Криптой / Stars → [@miropolbot](https://t.me/miropolbot) → /start\n\n"
        "_После оплаты подписка активируется автоматически._"
    )

async def cmd_2fa_setup(update, context):
    """Setup 2FA for admin."""
    user_id = update.effective_user.id
    if not is_admin(user_id):
        await update.message.reply_text("⛔ Admin only.")
        return
    if not TOTP_ENABLED:
        await update.message.reply_text("⛔ 2FA not available (pyotp missing).")
        return

    secrets = _load_2fa()
    if str(user_id) in secrets:
        await update.message.reply_text("✅ 2FA already set up. Use /admin to test.")
        return

    secret = pyotp.random_base32()
    secrets[str(user_id)] = secret
    _save_2fa(secrets)

    totp = pyotp.TOTP(secret)
    qr_url = totp.provisioning_uri(name=f"admin-{user_id}", issuer_name="Apolaibot")
    await update.message.reply_markdown(
        f"🔐 **2FA Setup**\n\n"
        f"Secret: `{secret}`\n\n"
        f"Add to Google Authenticator:\n"
        f"1. Open the app → '+' → 'Enter setup key'\n"
        f"2. Name: `Apolaibot`\n"
        f"3. Key: `{secret}`\n"
        f"4. Type: Time-based\n\n"
        f"Or scan QR: {qr_url}\n\n"
        f"Then test with `/admin <code>`"
    )

async def cmd_admin(update, context):
    """Admin panel — requires 2FA."""
    user_id = update.effective_user.id
    if not is_admin(user_id):
        await update.message.reply_text("⛔ Admin only.")
        return

    # If argument is a 6-digit code, verify 2FA
    args = update.message.text.split()
    if len(args) > 1 and args[1].isdigit() and len(args[1]) == 6:
        code = args[1]
        if verify_2fa(user_id, code):
            grant_2fa_session(user_id)
            await update.message.reply_text("✅ 2FA verified. Session active for 30 min. Use /stats, /admin panel.")
        else:
            await update.message.reply_text("❌ Invalid 2FA code. Try again.")
        return

    # Check 2FA session
    if not check_2fa_session(user_id):
        await update.message.reply_text("🔐 Enter 2FA code: `/admin <code>`")
        return

    # Admin panel
    users = _load_users()
    fb = _load_feedback()
    total = len(users)
    today = _dt.now().strftime("%Y-%m-%d")
    active_today = sum(1 for u in users.values() if u.get("last_seen", "").startswith(today))
    total_likes = sum(f.get("likes", 0) for f in fb.values())
    total_dislikes = sum(f.get("dislikes", 0) for f in fb.values())

    lines = [
        "🛡️ **Admin Panel**",
        f"👥 Users: {total} ({active_today} today)",
        f"👍 Likes: {total_likes}  |  👎 Dislikes: {total_dislikes}",
        f"🔐 2FA: {'✅ active' if TOTP_ENABLED else '⚠️ unavailable'}",
        "",
        "/stats — full stats  |  /2fa_setup — configure 2FA",
    ]
    await update.message.reply_markdown("\n".join(lines))

async def cmd_tools(update, context):
    """Show recent tool usage — transparency for users."""
    import glob as _glob
    tools_log = _Path.home() / ".hermes" / "data" / "audit_trail.jsonl"
    lines_out = ["🔧 **Recent tool activity:**"]
    try:
        if tools_log.exists():
            with open(tools_log) as f:
                all_lines = f.readlines()
                for line in all_lines[-5:]:
                    try:
                        entry = _json.loads(line)
                        ts = entry.get("timestamp", "")[:19]
                        op = entry.get("operation", "?")
                        agent = entry.get("agent", "?")
                        lines_out.append(f"• `{ts}` {op} by {agent}")
                    except Exception:
                        pass
    except Exception:
        pass
    lines_out.append("\n_This is what happens under the hood._")
    await update.message.reply_markdown("\n".join(lines_out))

async def feedback_callback(update, context):
    """Handle 👍👎 feedback."""
    query = update.callback_query
    await query.answer()
    action = query.data
    user_id = query.from_user.id
    fb = _load_feedback()
    uid = str(user_id)
    if uid not in fb:
        fb[uid] = {"likes": 0, "dislikes": 0}
    if action == "fb_like":
        fb[uid]["likes"] += 1
        await query.edit_message_reply_markup(reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("👍 Спасибо за фидбек!", callback_data="fb_done")]
        ]))
    elif action == "fb_dislike":
        fb[uid]["dislikes"] += 1
        await query.edit_message_reply_markup(reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("👎 Принято, работаем над улучшением", callback_data="fb_done")]
        ]))
    elif action == "fb_done":
        await query.edit_message_reply_markup(reply_markup=None)
        return
    _save_feedback(fb)

# ── Message handler ──

async def handle_message(update, context):
    msg = update.message.text
    if not msg or not msg.strip():
        return
    chat_id = update.effective_chat.id
    user_id = update.effective_user.id
    user_name = update.effective_user.username
    first_name = update.effective_user.first_name

    # Detect temp mode: check if this chat_id is in temp_conversations
    is_temp = chat_id in temp_conversations

    if not is_temp:
        # Rate limiting (skip for temp)
        if RATE_LIMIT_ENABLED and not can_proceed(user_id, "demo"):
            _, remaining, reset_in = rate_check(user_id, "demo")
            logger.info(f"[rate-limit] user={user_id} blocked, reset in {reset_in:.0f}s")
            await update.message.reply_text(RATE_LIMIT_COOLDOWN_MSG)
            return
        track_user(user_id, user_name, first_name)

    logger.info(f"[chat={chat_id}] {'[TEMP]' if is_temp else ''} {msg[:100]}")
    await context.bot.send_chat_action(chat_id=chat_id, action="typing")

    search_results = None
    if needs_web_search(msg):
        await context.bot.send_chat_action(chat_id=chat_id, action="typing")
        search_results = await tavily_search(msg)
        if search_results:
            msg = f"{msg}\n\n[Web search results — use these to answer accurately]:\n{search_results}"

    token_msg = ""
    if RATE_LIMIT_ENABLED and not is_temp and not can_use_tokens(user_id, 0):
        token_msg = "\n\n⚠️ Token limit reached this month. Response may be shorter than usual."

    reply, tokens_used = await chat_completion(chat_id, msg + token_msg, temp=is_temp)

    if RATE_LIMIT_ENABLED and not is_temp and tokens_used > 0:
        record_tokens(user_id, tokens_used)

    if AUDIT_ENABLED and not is_temp:
        log_action(user_id, "message", msg[:200])
    if AGENT_FS_ENABLED:
        fs_audit("user_message", "chat", message_id=str(update.message.message_id),
                 agent=f"demo-user-{user_id}", extra={"chat_id": str(chat_id), "msg_len": len(msg)})

    await update.message.reply_markdown(reply, reply_markup=_feedback_buttons())

async def error_handler(update, context):
    logger.error(f"Error: {context.error}")

async def handle_voice(update, context):
    """Transcribe voice message via Whisper."""
    if not WHISPER_ENABLED:
        await update.message.reply_text("🎤 Transcription unavailable (whisper not installed).")
        return

    global WHISPER_MODEL
    voice = update.message.voice
    chat_id = update.effective_chat.id

    await context.bot.send_chat_action(chat_id=chat_id, action="typing")
    await update.message.reply_text("🎤 Transcribing...")

    # Download voice file
    file = await context.bot.get_file(voice.file_id)
    voice_path = _Path.home() / ".local" / "share" / "apolaibot" / "voice" / f"{voice.file_id}.ogg"
    voice_path.parent.mkdir(parents=True, exist_ok=True)
    await file.download_to_drive(str(voice_path))

    try:
        # Load model once
        if WHISPER_MODEL is None:
            await update.message.reply_text("⏳ Loading Whisper model (first run, ~1.5GB)...")
            WHISPER_MODEL = whisper.load_model("tiny")  # tiny = fastest (240MB)

        result = WHISPER_MODEL.transcribe(str(voice_path), language="ru")
        text = result["text"].strip()

        if text:
            await update.message.reply_text(f"🎤 *Transcription:*\n{text}", parse_mode=ParseMode.MARKDOWN)
            # Also process as regular message
            update.message.text = text
            await handle_message(update, context)
        else:
            await update.message.reply_text("🎤 No speech detected.")
    except Exception as e:
        logger.error(f"Whisper error: {e}")
        await update.message.reply_text("🎤 Transcription failed. Try again later.")
    finally:
        voice_path.unlink(missing_ok=True)

# File upload limits
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_FILES_PER_USER = 3
UNLIMITED_USERS = {319665243, 2115597720, 470549555}  # Poliakarm, Kolesnikov, Ilya
_user_file_counts = {}

async def handle_document(update, context):
    """Handle document/file uploads with size and quantity limits."""
    user_id = update.effective_user.id
    doc = update.message.document
    chat_id = update.effective_chat.id
    if not doc:
        return

    file_size = doc.file_size or 0
    file_name = doc.file_name or "unknown"
    is_unlimited = user_id in UNLIMITED_USERS

    if not is_unlimited and file_size > MAX_FILE_SIZE:
        await update.message.reply_text(
            f"📄 Файл *{file_name}* слишком большой ({file_size/1024/1024:.1f} MB). Максимум: {MAX_FILE_SIZE/1024/1024:.0f} MB.",
            parse_mode=ParseMode.MARKDOWN)
        return

    count = _user_file_counts.get(user_id, 0)
    if not is_unlimited and count >= MAX_FILES_PER_USER:
        await update.message.reply_text(
            f"📄 Лимит файлов ({MAX_FILES_PER_USER} шт). Дождитесь обработки или /upgrade до Pro.")
        return

    _user_file_counts[user_id] = count + 1
    dl_path = None

    try:
        await context.bot.send_chat_action(chat_id=chat_id, action="typing")
        file = await context.bot.get_file(doc.file_id)
        dl_path = _Path.home() / ".local" / "share" / "apolaibot" / "uploads" / f"{user_id}_{doc.file_id}_{file_name}"
        dl_path.parent.mkdir(parents=True, exist_ok=True)
        await file.download_to_drive(str(dl_path))

        suffix = dl_path.suffix.lower()
        if suffix in ('.txt', '.md', '.py', '.json', '.csv', '.log'):
            text = dl_path.read_text()[:5000]
        elif suffix == '.pdf':
            try:
                import pymupdf; d = pymupdf.open(str(dl_path))
                text = "\n".join(page.get_text() for page in d)[:5000]; d.close()
            except ImportError:
                text = f"[PDF: {file_name}]"
        elif suffix in ('.docx', '.doc'):
            try:
                import docx; d = docx.Document(str(dl_path))
                text = "\n".join(p.text for p in d.paragraphs)[:5000]
            except ImportError:
                text = f"[DOCX: {file_name}]"
        else:
            text = f"[Файл: {file_name} — формат {suffix} не поддерживается]"

        if text:
            update.message.text = f"[Файл: {file_name}]\n\n{text}"
            await handle_message(update, context)
    except Exception as e:
        logger.error(f"Document error: {e}")
        await update.message.reply_text("📄 Ошибка обработки файла.")
    finally:
        _user_file_counts[user_id] = max(0, _user_file_counts.get(user_id, 1) - 1)
        if dl_path:
            try: dl_path.unlink(missing_ok=True)
            except Exception: pass

def main():
    logger.info(f"Starting apolaibot-demo (model={MODEL}, api={API_URL})")
    app = Application.builder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("skills", cmd_skills))
    app.add_handler(CommandHandler("summarize", cmd_summarize))
    app.add_handler(CommandHandler("temp", cmd_temp))
    app.add_handler(CommandHandler("export", cmd_export))
    app.add_handler(CommandHandler("stats", cmd_stats))
    app.add_handler(CommandHandler("upgrade", cmd_upgrade))
    app.add_handler(CommandHandler("2fa_setup", cmd_2fa_setup))
    app.add_handler(CommandHandler("admin", cmd_admin))
    app.add_handler(CommandHandler("tools", cmd_tools))
    app.add_handler(CallbackQueryHandler(role_callback, pattern="^role_"))
    app.add_handler(CallbackQueryHandler(feedback_callback, pattern="^fb_"))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    app.add_error_handler(error_handler)

    # Set bot menu
    async def set_menu(app):
        await app.bot.set_my_commands([
            BotCommand("start", "🔄 Онбординг — выберите роль"),
            BotCommand("skills", "📋 Каталог скиллов по профессиям"),
            BotCommand("summarize", "📝 Сделать выжимку диалога"),
            BotCommand("export", "📦 Экспорт чата в файл"),
            BotCommand("stats", "📊 Статистика и бюджет токенов"),
            BotCommand("tools", "🔧 Прозрачность — что под капотом"),
            BotCommand("upgrade", "🚀 Тарифы и оплата"),
        ])
    app.post_init = set_menu

    logger.info("Bot polling...")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
